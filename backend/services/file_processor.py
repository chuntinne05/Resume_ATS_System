import os
import tempfile
from typing import Dict, Any, Optional
import PyPDF2
import pdfplumber
from docx import Document
import boto3
from PIL import Image
import pytesseract
import logging

class FileProcessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.textract_client = boto3.client(
            'textract',
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
            region_name=os.getenv('AWS_REGION', 'ap-southeast-1')
        )
    
    def extract_text_from_file(self, file_content: bytes, filename: str, s3_key: str) -> Dict[str, Any]:
        file_extension = os.path.splitext(filename)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_content)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff']:
                return self._extract_from_image(file_content, s3_key)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_extension}'
                }
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        text = ""
        page_count = 0
        
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                with pdfplumber.open(temp_file.name) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            if not text.strip():
                with tempfile.NamedTemporaryFile() as temp_file:
                    temp_file.write(file_content)
                    temp_file.flush()
                    
                    with open(temp_file.name, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        page_count = len(pdf_reader.pages)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'page_count': page_count,
                'extraction_method': 'PDF',
                'word_count': len(text.split())
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'PDF extraction failed: {str(e)}'
            }
    
    def _extract_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                doc = Document(temp_file.name)
                text = ""
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\t"
                        text += "\n"
                
                return {
                    'success': True,
                    'text': text.strip(),
                    'page_count': 1,
                    'extraction_method': 'DOCX',
                    'word_count': len(text.split())
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX extraction failed: {str(e)}'
            }
    
    def _extract_from_image(self, file_content: bytes, s3_key: str) -> Dict[str, Any]:
        try:
            # Try AWS Textract first
            textract_result = self._extract_with_textract(s3_key)
            if textract_result['success']:
                return textract_result
            
            # Fallback to Tesseract
            return self._extract_with_tesseract(file_content)
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Image extraction failed: {str(e)}'
            }
    
    def _extract_with_textract(self, s3_key: str) -> Dict[str, Any]:
        try:
            response = self.textract_client.detect_document_text(
                Document={
                    'S3Object': {
                        'Bucket': os.getenv('S3_BUCKET_NAME'),
                        'Name': s3_key
                    }
                }
            )
            
            text = ""
            for block in response['Blocks']:
                if block['BlockType'] == 'LINE':
                    text += block['Text'] + "\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'page_count': 1,
                'extraction_method': 'AWS_Textract',
                'word_count': len(text.split()),
                'confidence': response.get('DetectDocumentTextModelVersion', 'N/A')
            }
            
        except Exception as e:
            self.logger.warning(f"Textract failed: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def _extract_with_tesseract(self, file_content: bytes) -> Dict[str, Any]:
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                image = Image.open(temp_file.name)
                text = pytesseract.image_to_string(image, lang='eng+vie')
                
                return {
                    'success': True,
                    'text': text.strip(),
                    'page_count': 1,
                    'extraction_method': 'Tesseract_OCR',
                    'word_count': len(text.split())
                }
                
        except Exception as e:
            self.logger.error(f"Tesseract failed: {str(e)}")
            return {
                'success': False,
                'error': f'Tesseract extraction failed: {str(e)}'
            }

# Singleton instance
file_processor = FileProcessor()